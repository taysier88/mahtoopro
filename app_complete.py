import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, white, black
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display
import io
import base64
import plotly.express as px
import plotly.graph_objects as go

st.set_page_config(page_title="محتوى برو - نظام إدارة المشاريع", page_icon="🚀", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700&display=swap');
.main { direction: rtl; text-align: right; font-family: 'Cairo', sans-serif; }
.stButton>button { width: 100%; border-radius: 8px; font-weight: bold; }
.stTextInput>div>div>input, .stTextArea>div>div>textarea { direction: rtl; text-align: right; }
.stSelectbox>div>div>div { direction: rtl; text-align: right; }
.metric-card {
    background: linear-gradient(135deg, #1a5f7a 0%, #159895 100%);
    color: white; padding: 20px; border-radius: 15px;
    text-align: center; box-shadow: 0 4px 6px rgba(0,0,0,0.1);
}
.metric-value { font-size: 32px; font-weight: bold; }
.metric-label { font-size: 14px; opacity: 0.9; }
.card {
    background: #f8f9fa; border-radius: 12px; padding: 15px;
    margin: 8px 0; border-right: 4px solid #1a5f7a;
}
.header-gradient {
    background: linear-gradient(135deg, #1a5f7a 0%, #159895 100%);
    color: white; padding: 25px; border-radius: 15px; margin-bottom: 25px;
}
</style>
""", unsafe_allow_html=True)

DATA_DIR = "data"
FILES = {
    "clients": os.path.join(DATA_DIR, "clients.json"),
    "projects": os.path.join(DATA_DIR, "projects.json"),
    "tasks": os.path.join(DATA_DIR, "tasks.json"),
    "content": os.path.join(DATA_DIR, "content_calendar.json"),
    "payments": os.path.join(DATA_DIR, "payments.json"),
    "settings": os.path.join(DATA_DIR, "settings.json"),
}

def init_data():
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    for key, path in FILES.items():
        if not os.path.exists(path):
            with open(path, 'w', encoding='utf-8') as f:
                if key == "settings":
                    json.dump({"company_name": "محتوى برو", "phone": "01XXXXXXXXX", "email": "info@mahtoobro.com",
                               "basic_price": "2,000 - 3,500", "standard_price": "4,500 - 8,000", "pro_price": "9,500 - 18,000"}, f, ensure_ascii=False)
                else:
                    json.dump([], f)

def load_data(key):
    try:
        with open(FILES[key], 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return []

def save_data(key, data):
    with open(FILES[key], 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def get_next_id(data_list):
    if not data_list:
        return 1
    return max(item.get("id", 0) for item in data_list) + 1

try:
    init_data()
except:
    pass

settings = load_data("settings")

def ar(text):
    try:
        reshaped = arabic_reshaper.reshape(str(text))
        return get_display(reshaped)
    except:
        return str(text)

try:
    pdfmetrics.registerFont(TTFont('Amiri', 'Amiri-Regular.ttf'))
    ARABIC_FONT = 'Amiri'
except:
    ARABIC_FONT = 'Helvetica'

def generate_contract_doc(client, project):
    doc = Document()
    def add_para(text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(size)
        run.bold = bold
        return p
    def add_head(text, level=1):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in h.runs:
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        return h
    add_head('عقد خدمات التسويق الرقمي', level=0)
    add_para(f'{settings.get("company_name", "محتوى برو")} - إدارة حسابات السوشيال ميديا', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('─────────────────────────────────────────', align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para('')
    add_head('أطراف العقد', level=1)
    add_para(f'الطرف الأول: مقدم الخدمة ({settings.get("company_name", "محتوى برو")})')
    add_para(f'الطرف الثاني: {client.get("name", "")}')
    add_para(f'العنوان: {client.get("address", "")}')
    add_para(f'الهاتف: {client.get("phone", "")}')
    add_para('')
    add_head('تفاصيل المشروع', level=1)
    add_para(f'اسم المشروع: {project.get("name", "")}')
    add_para(f'الباقة: {project.get("package", "")}')
    add_para(f'السعر: {project.get("price", "")} ج.م/شهر')
    add_para(f'تاريخ البدء: {project.get("start_date", "")}')
    add_para(f'تاريخ الانتهاء: {project.get("end_date", "")}')
    add_para(f'الحالة: {project.get("status", "")}')
    add_para('')
    add_head('نطاق العمل', level=1)
    add_para(project.get("scope", "إدارة حسابات السوشيال ميديا وإنشاء محتوى"))
    add_para('')
    add_head('شروط الدفع', level=1)
    add_para('• الدفع شهرياً مقدماً')
    add_para('• خصم 10% للاشتراك 3 أشهر مقدماً')
    add_para('• يتم الدفع عبر: فودافون كاش / إنستا باي / تحويل بنكي')
    add_para('')
    add_head('التوقيعات', level=1)
    add_para(f'الطرف الأول: {settings.get("company_name", "")}          الطرف الثاني: {client.get("name", "")}')
    add_para('التوقيع: ___________          التوقيع: ___________')
    add_para(f'التاريخ: {datetime.now().strftime("%Y-%m-%d")}')
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_quotation_pdf(client_name, package, price, notes=""):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    primary = HexColor('#1a5f7a')
    accent = HexColor('#159895')
    c.setFillColor(primary)
    c.rect(0, height - 3.5*cm, width, 3.5*cm, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont(ARABIC_FONT, 24)
    c.drawRightString(width - 2*cm, height - 2*cm, ar(settings.get("company_name", "محتوى برو")))
    c.setFont(ARABIC_FONT, 12)
    c.drawRightString(width - 2*cm, height - 2.8*cm, ar('عرض أسعار'))
    c.setFillColor(black)
    c.setFont(ARABIC_FONT, 11)
    y = height - 5*cm
    c.drawRightString(width - 2*cm, y, ar(f'العميل: {client_name}'))
    c.drawRightString(width - 2*cm, y - 0.7*cm, ar(f'الباقة: {package}'))
    c.drawRightString(width - 2*cm, y - 1.4*cm, ar(f'السعر: {price} ج.م/شهر'))
    c.drawRightString(width - 2*cm, y - 2.1*cm, ar(f'التاريخ: {datetime.now().strftime("%Y-%m-%d")}'))
    c.setFillColor(HexColor('#f5f5f5'))
    c.roundRect(2*cm, height - 10*cm, width - 4*cm, 3.5*cm, 10, fill=1, stroke=0)
    c.setFillColor(primary)
    c.setFont(ARABIC_FONT, 13)
    c.drawRightString(width - 3*cm, height - 7.5*cm, ar('تفاصيل الباقة'))
    details = {
        'الأساسية': ['12 منشور/شهر', '3 ريلز/أسبوع', 'تصميم صور + كتابة محتوى', 'تقرير شهري'],
        'القياسية': ['20 منشور/شهر', '4 ريلز + فيديو', 'تفاعل يومي', 'تقرير + مقترحات نمو'],
        'الاحترافية': ['30 منشور/شهر', '8 ريلز + 4 فيديوهات', 'تصوير دوري', 'إعلانات مدفوعة']
    }
    c.setFillColor(black)
    c.setFont(ARABIC_FONT, 10)
    fy = height - 8.2*cm
    for feat in details.get(package, []):
        c.drawRightString(width - 3*cm, fy, ar('• ' + feat))
        fy -= 0.5*cm
    if notes:
        c.setFillColor(accent)
        c.setFont(ARABIC_FONT, 10)
        c.drawRightString(width - 2*cm, height - 11.5*cm, ar('ملاحظات:'))
        c.setFillColor(black)
        c.drawRightString(width - 2*cm, height - 12.1*cm, ar(notes))
    c.setFillColor(primary)
    c.rect(0, 0, width, 1.5*cm, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont(ARABIC_FONT, 9)
    c.drawCentredString(width/2, 0.7*cm, ar(f'{settings.get("company_name", "")} - {settings.get("phone", "")}'))
    c.save()
    buffer.seek(0)
    return buffer

def get_download_link(buffer, filename, label):
    b64 = base64.b64encode(buffer.read()).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}"><button style="background:#1a5f7a;color:white;padding:8px 16px;border:none;border-radius:8px;cursor:pointer;">{label}</button></a>'

with st.sidebar:
    st.markdown(f"## 🚀 {settings.get('company_name', 'محتوى برو')}")
    st.markdown("---")
    menu = st.radio("القائمة", [
        "🏠 لوحة التحكم", "👥 العملاء", "📋 المشاريع", "✅ المهام",
        "📝 التقويم المحتوي", "💰 المالية", "📊 التقارير", "📄 المستندات", "⚙️ الإعدادات"
    ], label_visibility="collapsed")
    st.markdown("---")
    st.caption(f"📅 {datetime.now().strftime('%Y-%m-%d')}")

if menu == "🏠 لوحة التحكم":
    st.markdown('<div class="header-gradient"><h1>🏠 لوحة التحكم</h1><p>نظرة عامة على أداء عملك</p></div>', unsafe_allow_html=True)
    clients = load_data("clients")
    projects = load_data("projects")
    tasks = load_data("tasks")
    payments = load_data("payments")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        active_projects = len([p for p in projects if p.get("status") == "نشط"])
        st.markdown(f'<div class="metric-card"><div class="metric-value">{active_projects}</div><div class="metric-label">مشاريع نشطة</div></div>', unsafe_allow_html=True)
    with col2:
        total_clients = len(clients)
        st.markdown(f'<div class="metric-card"><div class="metric-value">{total_clients}</div><div class="metric-label">إجمالي العملاء</div></div>', unsafe_allow_html=True)
    with col3:
        pending_tasks = len([t for t in tasks if t.get("status") == "قيد التنفيذ"])
        st.markdown(f'<div class="metric-card"><div class="metric-value">{pending_tasks}</div><div class="metric-label">مهام معلقة</div></div>', unsafe_allow_html=True)
    with col4:
        total_revenue = sum(float(p.get("amount", 0)) for p in payments if p.get("status") == "مدفوع")
        st.markdown(f'<div class="metric-card"><div class="metric-value">{total_revenue:,.0f}</div><div class="metric-label">إيرادات (ج.م)</div></div>', unsafe_allow_html=True)

    st.markdown("---")
    col_left, col_right = st.columns(2)
    with col_left:
        st.markdown("### 📋 آخر المشاريع")
        if projects:
            recent = sorted(projects, key=lambda x: x.get("start_date", ""), reverse=True)[:5]
            for p in recent:
                status_color = "status-active" if p.get("status") == "نشط" else "status-pending" if p.get("status") == "قيد الانتظار" else "status-overdue"
                st.markdown(f'<div class="card"><b>{p.get("name", "")}</b> - <span style="color:#28a745;font-weight:bold;">{p.get("status", "")}</span><br><small>العميل: {p.get("client_name", "")} | السعر: {p.get("price", "")} ج.م</small></div>', unsafe_allow_html=True)
        else:
            st.info("لا توجد مشاريع بعد. أضف مشروعاً من قسم المشاريع.")
    with col_right:
        st.markdown("### ✅ المهام القادمة")
        if tasks:
            upcoming = [t for t in tasks if t.get("status") != "مكتمل"]
            upcoming = sorted(upcoming, key=lambda x: x.get("due_date", "9999-12-31"))[:5]
            for t in upcoming:
                overdue = datetime.strptime(t.get("due_date", "9999-12-31"), "%Y-%m-%d").date() < datetime.now().date()
                status_text = "متأخر" if overdue else t.get("status", "")
                color = "#dc3545" if overdue else "#ffc107"
                st.markdown(f'<div class="card"><b>{t.get("title", "")}</b> - <span style="color:{color};font-weight:bold;">{status_text}</span><br><small>تاريخ التسليم: {t.get("due_date", "")} | المشروع: {t.get("project_name", "")}</small></div>', unsafe_allow_html=True)
        else:
            st.info("لا توجد مهام معلقة.")

    st.markdown("---")
    st.markdown("### 📈 الإيرادات الشهرية")
    if payments:
        df_pay = pd.DataFrame(payments)
        df_pay["date"] = pd.to_datetime(df_pay["date"])
        df_pay["month"] = df_pay["date"].dt.to_period("M").astype(str)
        monthly = df_pay[df_pay["status"] == "مدفوع"].groupby("month")["amount"].sum().reset_index()
        if not monthly.empty:
            fig = px.bar(monthly, x="month", y="amount", labels={"month": "الشهر", "amount": "المبلغ (ج.م)"}, color_discrete_sequence=["#1a5f7a"])
            fig.update_layout(xaxis_title="الشهر", yaxis_title="المبلغ (ج.م)", showlegend=False)
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("لا توجد دفعات مسجلة بعد.")
    else:
        st.info("سجل دفعاتك في قسم المالية لرؤية التحليلات.")

elif menu == "👥 العملاء":
    st.markdown('<div class="header-gradient"><h1>👥 إدارة العملاء</h1></div>', unsafe_allow_html=True)
    clients = load_data("clients")
    tab1, tab2 = st.tabs(["📋 قائمة العملاء", "➕ إضافة عميل جديد"])
    with tab1:
        if clients:
            df = pd.DataFrame(clients)
            st.dataframe(df[["id", "name", "phone", "email", "address", "category", "notes"]], use_container_width=True, hide_index=True)
            col1, col2 = st.columns(2)
            with col1:
                if st.button("🗑️ حذف الكل", type="secondary"):
                    save_data("clients", [])
                    st.rerun()
            with col2:
                csv = df.to_csv(index=False).encode('utf-8-sig')
                st.download_button("📥 تصدير CSV", csv, "clients.csv", "text/csv")
        else:
            st.info("لا يوجد عملاء مسجلين. أضف عميلاً جديداً.")
    with tab2:
        with st.form("client_form"):
            col1, col2 = st.columns(2)
            with col1:
                name = st.text_input("اسم العميل / المطعم *", placeholder="مثال: مطعم الأصالة")
                phone = st.text_input("رقم الهاتف *", placeholder="01XXXXXXXXX")
                email = st.text_input("البريد الإلكتروني", placeholder="example@email.com")
            with col2:
                address = st.text_input("العنوان", placeholder="القاهرة - مصر الجديدة")
                category = st.selectbox("التصنيف", ["مطعم", "كافيه", "مخبز", "حلويات", "أخرى"])
                social = st.text_input("رابط السوشيال ميديا", placeholder="رابط انستغرام أو فيسبوك")
            notes = st.text_area("ملاحظات", placeholder="أي ملاحظات عن العميل...")
            submitted = st.form_submit_button("💾 حفظ العميل", type="primary")
            if submitted:
                if not name or not phone:
                    st.error("❌ الاسم والهاتف مطلوبان")
                else:
                    new_client = {
                        "id": get_next_id(clients), "name": name, "phone": phone,
                        "email": email, "address": address, "category": category,
                        "social": social, "notes": notes,
                        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M")
                    }
                    clients.append(new_client)
                    save_data("clients", clients)
                    st.success(f"✅ تم إضافة العميل: {name}")
                    st.rerun()

elif menu == "📋 المشاريع":
    st.markdown('<div class="header-gradient"><h1>📋 إدارة المشاريع</h1></div>', unsafe_allow_html=True)
    clients = load_data("clients")
    projects = load_data("projects")
    tab1, tab2 = st.tabs(["📋 المشاريع الحالية", "➕ مشروع جديد"])
    with tab1:
        if projects:
            for p in projects:
                status_emoji = {"نشط": "🟢", "قيد الانتظار": "🟡", "منتهي": "🔵", "ملغي": "🔴"}.get(p.get("status"), "⚪")
                with st.expander(f"{status_emoji} {p.get('name', '')} - {p.get('client_name', '')} ({p.get('price', '')} ج.م)"):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.write(f"**الباقة:** {p.get('package', '')}")
                        st.write(f"**تاريخ البدء:** {p.get('start_date', '')}")
                        st.write(f"**تاريخ الانتهاء:** {p.get('end_date', '')}")
                    with col2:
                        st.write(f"**الحالة:** {p.get('status', '')}")
                        st.write(f"**نطاق العمل:** {p.get('scope', '')}")
                    col_a, col_b, col_c = st.columns(3)
                    with col_a:
                        new_status = st.selectbox("تغيير الحالة", ["نشط", "قيد الانتظار", "منتهي", "ملغي"], 
                                                  index=["نشط", "قيد الانتظار", "منتهي", "ملغي"].index(p.get("status", "نشط")), key=f"status_{p['id']}")
                        if new_status != p.get("status"):
                            p["status"] = new_status
                            save_data("projects", projects)
                            st.rerun()
                    with col_b:
                        if st.button("📄 عقد", key=f"contract_{p['id']}"):
                            client = next((c for c in clients if c["id"] == p.get("client_id")), {})
                            buf = generate_contract_doc(client, p)
                            st.markdown(get_download_link(buf, f"عقد_{p['name']}.docx", "📥 تحميل العقد"), unsafe_allow_html=True)
                    with col_c:
                        if st.button("🗑️ حذف", key=f"del_proj_{p['id']}"):
                            projects = [pr for pr in projects if pr["id"] != p["id"]]
                            save_data("projects", projects)
                            st.rerun()
        else:
            st.info("لا توجد مشاريع. أضف مشروعاً جديداً.")
    with tab2:
        if not clients:
            st.warning("⚠️ أضف عميلاً أولاً من قسم العملاء")
        else:
            with st.form("project_form"):
                col1, col2 = st.columns(2)
                with col1:
                    client_sel = st.selectbox("العميل", options=clients, format_func=lambda x: x["name"])
                    name = st.text_input("اسم المشروع", placeholder="مثال: إدارة حساب مطعم الأصالة")
                    package = st.selectbox("الباقة", ["الأساسية", "القياسية", "الاحترافية"])
                with col2:
                    price_map = {"الأساسية": settings.get("basic_price", "2,000 - 3,500"),
                                 "القياسية": settings.get("standard_price", "4,500 - 8,000"),
                                 "الاحترافية": settings.get("pro_price", "9,500 - 18,000")}
                    price = st.text_input("السعر (ج.م/شهر)", value=price_map.get(package, ""))
                    start_date = st.date_input("تاريخ البدء", value=datetime.now())
                    duration_months = st.number_input("المدة (شهور)", min_value=1, max_value=24, value=1)
                scope = st.text_area("نطاق العمل", placeholder="وصف تفصيلي للخدمات المقدمة...")
                submitted = st.form_submit_button("💾 إنشاء المشروع", type="primary")
                if submitted:
                    if not name:
                        st.error("❌ اسم المشروع مطلوب")
                    else:
                        end_date = start_date + timedelta(days=30*duration_months)
                        new_project = {
                            "id": get_next_id(projects), "client_id": client_sel["id"],
                            "client_name": client_sel["name"], "name": name,
                            "package": package, "price": price,
                            "start_date": start_date.strftime("%Y-%m-%d"),
                            "end_date": end_date.strftime("%Y-%m-%d"),
                            "status": "قيد الانتظار", "scope": scope,
                            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M")
                        }
                        projects.append(new_project)
                        save_data("projects", projects)
                        st.success(f"✅ تم إنشاء المشروع: {name}")
                        st.rerun()

elif menu == "✅ المهام":
    st.markdown('<div class="header-gradient"><h1>✅ إدارة المهام</h1></div>', unsafe_allow_html=True)
    tasks = load_data("tasks")
    projects = load_data("projects")
    col1, col2 = st.columns([2, 1])
    with col1:
        st.markdown("### 📋 قائمة المهام")
        if tasks:
            filter_status = st.selectbox("تصفية حسب الحالة", ["الكل", "قيد التنفيذ", "مكتمل", "متأخر"])
            for t in tasks:
                is_overdue = datetime.strptime(t.get("due_date", "9999-12-31"), "%Y-%m-%d").date() < datetime.now().date() and t.get("status") != "مكتمل"
                if filter_status == "متأخر" and not is_overdue:
                    continue
                if filter_status != "الكل" and filter_status != "متأخر" and t.get("status") != filter_status:
                    continue
                if filter_status == "متأخر" and is_overdue:
                    pass
                elif filter_status == "الكل":
                    pass
                elif t.get("status") != filter_status:
                    continue
                status_emoji = {"قيد التنفيذ": "⏳", "مكتمل": "✅"}.get(t.get("status"), "⚪")
                if is_overdue:
                    status_emoji = "🔴"
                with st.container():
                    cols = st.columns([4, 2, 1, 1])
                    with cols[0]:
                        st.write(f"{status_emoji} **{t.get('title', '')}**")
                        st.caption(f"المشروع: {t.get('project_name', '')} | تسليم: {t.get('due_date', '')}")
                    with cols[1]:
                        priority_color = {"عالية": "🔴", "متوسطة": "🟡", "منخفضة": "🟢"}.get(t.get("priority", ""), "⚪")
                        st.write(f"{priority_color} {t.get('priority', '')}")
                    with cols[2]:
                        if t.get("status") != "مكتمل":
                            if st.button("✅ إنجاز", key=f"done_{t['id']}"):
                                t["status"] = "مكتمل"
                                t["completed_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                                save_data("tasks", tasks)
                                st.rerun()
                    with cols[3]:
                        if st.button("🗑️", key=f"del_task_{t['id']}"):
                            tasks = [tk for tk in tasks if tk["id"] != t["id"]]
                            save_data("tasks", tasks)
                            st.rerun()
                    st.markdown("---")
        else:
            st.info("لا توجد مهام. أضف مهمة جديدة.")
    with col2:
        st.markdown("### ➕ مهمة جديدة")
        with st.form("task_form"):
            title = st.text_input("عنوان المهمة", placeholder="مثال: تصميم منشورات الأسبوع الأول")
            project_options = [{"id": 0, "name": "عام"}] + projects
            project_sel = st.selectbox("المشروع", options=project_options, format_func=lambda x: x["name"])
            due_date = st.date_input("تاريخ التسليم", value=datetime.now() + timedelta(days=3))
            priority = st.selectbox("الأولوية", ["عالية", "متوسطة", "منخفضة"])
            notes = st.text_area("ملاحظات", placeholder="تفاصيل إضافية...")
            submitted = st.form_submit_button("💾 إضافة", type="primary")
            if submitted:
                if not title:
                    st.error("❌ عنوان المهمة مطلوب")
                else:
                    new_task = {
                        "id": get_next_id(tasks), "title": title,
                        "project_id": project_sel.get("id", 0),
                        "project_name": project_sel.get("name", "عام"),
                        "due_date": due_date.strftime("%Y-%m-%d"),
                        "priority": priority, "status": "قيد التنفيذ",
                        "notes": notes,
                        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M")
                    }
                    tasks.append(new_task)
                    save_data("tasks", tasks)
                    st.success("✅ تم إضافة المهمة")
                    st.rerun()

elif menu == "📝 التقويم المحتوي":
    st.markdown('<div class="header-gradient"><h1>📝 التقويم المحتوي</h1></div>', unsafe_allow_html=True)
    content = load_data("content")
    projects = load_data("projects")
    tab1, tab2 = st.tabs(["📅 عرض التقويم", "➕ إضافة محتوى"])
    with tab1:
        if content:
            project_filter = st.selectbox("تصفية حسب المشروع", ["الكل"] + [p["name"] for p in projects])
            df_content = pd.DataFrame(content)
            if project_filter != "الكل":
                df_content = df_content[df_content["project_name"] == project_filter]
            df_content["publish_date"] = pd.to_datetime(df_content["publish_date"])
            df_content = df_content.sort_values("publish_date")
            for _, row in df_content.iterrows():
                status_emoji = {"منشور": "✅", "جاهز": "📦", "قيد التصميم": "🎨", "مخطط": "📋"}.get(row["status"], "⚪")
                with st.container():
                    cols = st.columns([1, 3, 2, 1, 1])
                    cols[0].write(f"**{row['publish_date'].strftime('%Y-%m-%d')}**")
                    cols[1].write(f"{status_emoji} {row['title']}")
                    cols[2].write(f"📱 {row['platform']} | {row['type']}")
                    cols[3].write(f"📋 {row['project_name']}")
                    if cols[4].button("🗑️", key=f"del_content_{row['id']}"):
                        content = [c for c in content if c["id"] != row["id"]]
                        save_data("content", content)
                        st.rerun()
                    st.markdown("---")
        else:
            st.info("لا يوجد محتوى مخطط. أضف منشورات من التبويب التالي.")
    with tab2:
        with st.form("content_form"):
            col1, col2 = st.columns(2)
            with col1:
                project_sel = st.selectbox("المشروع", options=projects, format_func=lambda x: x["name"])
                title = st.text_input("عنوان المنشور", placeholder="مثال: عرض خاص نهاية الأسبوع")
                content_type = st.selectbox("نوع المحتوى", ["صورة", "ريلز", "فيديو", "قصة", "كاروسيل", "مقال"])
            with col2:
                platform = st.selectbox("المنصة", ["انستغرام", "فيسبوك", "تيك توك", "تويتر", "لينكدإن"])
                publish_date = st.date_input("تاريخ النشر")
                status = st.selectbox("الحالة", ["مخطط", "قيد التصميم", "جاهز", "منشور"])
            caption = st.text_area("نص المنشور / الكابشن", placeholder="اكتب نص المنشور هنا...")
            submitted = st.form_submit_button("💾 حفظ المحتوى", type="primary")
            if submitted:
                if not title:
                    st.error("❌ عنوان المنشور مطلوب")
                else:
                    new_content = {
                        "id": get_next_id(content), "project_id": project_sel["id"],
                        "project_name": project_sel["name"], "title": title,
                        "type": content_type, "platform": platform,
                        "publish_date": publish_date.strftime("%Y-%m-%d"),
                        "status": status, "caption": caption,
                        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M")
                    }
                    content.append(new_content)
                    save_data("content", content)
                    st.success("✅ تم إضافة المحتوى")
                    st.rerun()

elif menu == "💰 المالية":
    st.markdown('<div class="header-gradient"><h1>💰 إدارة المالية</h1></div>', unsafe_allow_html=True)
    payments = load_data("payments")
    projects = load_data("projects")
    total_paid = sum(float(p.get("amount", 0)) for p in payments if p.get("status") == "مدفوع")
    total_pending = sum(float(p.get("amount", 0)) for p in payments if p.get("status") == "معلق")
    total_overdue = sum(float(p.get("amount", 0)) for p in payments if p.get("status") == "متأخر")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown(f'<div class="metric-card"><div class="metric-value">{total_paid:,.0f}</div><div class="metric-label">💰 مدفوع</div></div>', unsafe_allow_html=True)
    with col2:
        st.markdown(f'<div class="metric-card" style="background:linear-gradient(135deg, #ffc107 0%, #ff9800 100%);"><div class="metric-value">{total_pending:,.0f}</div><div class="metric-label">⏳ معلق</div></div>', unsafe_allow_html=True)
    with col3:
        st.markdown(f'<div class="metric-card" style="background:linear-gradient(135deg, #dc3545 0%, #c82333 100%);"><div class="metric-value">{total_overdue:,.0f}</div><div class="metric-label">🔴 متأخر</div></div>', unsafe_allow_html=True)
    st.markdown("---")
    tab1, tab2 = st.tabs(["📋 سجل الدفعات", "➕ تسجيل دفعة"])
    with tab1:
        if payments:
            df_pay = pd.DataFrame(payments)
            st.dataframe(df_pay[["id", "project_name", "amount", "status", "date", "method", "notes"]], use_container_width=True, hide_index=True)
            st.markdown("### ✏️ تحديث حالة الدفعة")
            pay_options = [f"{p['project_name']} - {p['amount']} ج.م ({p['status']})" for p in payments]
            pay_idx = st.selectbox("اختر الدفعة", range(len(payments)), format_func=lambda i: pay_options[i])
            new_pay_status = st.selectbox("الحالة الجديدة", ["مدفوع", "معلق", "متأخر"])
            if st.button("💾 تحديث"):
                payments[pay_idx]["status"] = new_pay_status
                if new_pay_status == "مدفوع":
                    payments[pay_idx]["paid_date"] = datetime.now().strftime("%Y-%m-%d")
                save_data("payments", payments)
                st.rerun()
        else:
            st.info("لا توجد دفعات مسجلة.")
    with tab2:
        with st.form("payment_form"):
            col1, col2 = st.columns(2)
            with col1:
                project_sel = st.selectbox("المشروع", options=projects, format_func=lambda x: x["name"])
                amount = st.number_input("المبلغ (ج.م)", min_value=0, step=100)
                pay_status = st.selectbox("الحالة", ["معلق", "مدفوع", "متأخر"])
            with col2:
                pay_date = st.date_input("تاريخ الاستحقاق")
                method = st.selectbox("طريقة الدفع", ["فودافون كاش", "إنستا باي", "تحويل بنكي", "كاش"])
            notes = st.text_area("ملاحظات", placeholder="رقم العملية أو تفاصيل إضافية...")
            submitted = st.form_submit_button("💾 تسجيل الدفعة", type="primary")
            if submitted:
                if amount <= 0:
                    st.error("❌ المبلغ يجب أن يكون أكبر من صفر")
                else:
                    new_payment = {
                        "id": get_next_id(payments), "project_id": project_sel["id"],
                        "project_name": project_sel["name"], "amount": amount,
                        "status": pay_status, "date": pay_date.strftime("%Y-%m-%d"),
                        "method": method, "notes": notes,
                        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M")
                    }
                    if pay_status == "مدفوع":
                        new_payment["paid_date"] = datetime.now().strftime("%Y-%m-%d")
                    payments.append(new_payment)
                    save_data("payments", payments)
                    st.success("✅ تم تسجيل الدفعة")
                    st.rerun()

elif menu == "📊 التقارير":
    st.markdown('<div class="header-gradient"><h1>📊 التقارير والتحليلات</h1></div>', unsafe_allow_html=True)
    clients = load_data("clients")
    projects = load_data("projects")
    tasks = load_data("tasks")
    payments = load_data("payments")
    content = load_data("content")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("### 📈 توزيع المشاريع حسب الحالة")
        if projects:
            df_proj = pd.DataFrame(projects)
            status_counts = df_proj["status"].value_counts().reset_index()
            status_counts.columns = ["الحالة", "العدد"]
            fig = px.pie(status_counts, values="العدد", names="الحالة", color_discrete_sequence=["#1a5f7a", "#57c5b6", "#159895", "#dc3545"])
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("لا توجد بيانات كافية")
    with col2:
        st.markdown("### 📊 توزيع العملاء حسب التصنيف")
        if clients:
            df_cli = pd.DataFrame(clients)
            cat_counts = df_cli["category"].value_counts().reset_index()
            cat_counts.columns = ["التصنيف", "العدد"]
            fig = px.bar(cat_counts, x="التصنيف", y="العدد", color_discrete_sequence=["#159895"])
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("لا توجد بيانات كافية")
    st.markdown("---")
    col3, col4 = st.columns(2)
    with col3:
        st.markdown("### ✅ إنجاز المهام")
        if tasks:
            df_tasks = pd.DataFrame(tasks)
            task_status = df_tasks["status"].value_counts().reset_index()
            task_status.columns = ["الحالة", "العدد"]
            fig = px.bar(task_status, x="الحالة", y="العدد", color="الحالة", color_discrete_map={"مكتمل": "#28a745", "قيد التنفيذ": "#ffc107"})
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("لا توجد مهام مسجلة")
    with col4:
        st.markdown("### 📝 حالة المحتوى")
        if content:
            df_cont = pd.DataFrame(content)
            cont_status = df_cont["status"].value_counts().reset_index()
            cont_status.columns = ["الحالة", "العدد"]
            fig = px.pie(cont_status, values="العدد", names="الحالة", color_discrete_sequence=["#1a5f7a", "#57c5b6", "#159895", "#28a745"])
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("لا يوجد محتوى مسجل")
    st.markdown("---")
    st.markdown("### 💰 تقرير الإيرادات الشهرية")
    if payments:
        df_pay = pd.DataFrame(payments)
        df_pay["date"] = pd.to_datetime(df_pay["date"])
        df_pay["month"] = df_pay["date"].dt.to_period("M").astype(str)
        monthly = df_pay.groupby(["month", "status"])["amount"].sum().reset_index()
        fig = px.bar(monthly, x="month", y="amount", color="status", labels={"month": "الشهر", "amount": "المبلغ (ج.م)", "status": "الحالة"}, color_discrete_map={"مدفوع": "#28a745", "معلق": "#ffc107", "متأخر": "#dc3545"})
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("سجل الدفعات في قسم المالية لرؤية التحليلات")

elif menu == "📄 المستندات":
    st.markdown('<div class="header-gradient"><h1>📄 مولد المستندات</h1></div>', unsafe_allow_html=True)
    clients = load_data("clients")
    tab1, tab2 = st.tabs(["📋 عقد جديد", "💰 عرض أسعار"])
    with tab1:
        if not clients:
            st.warning("⚠️ أضف عميلاً أولاً من قسم العملاء")
        else:
            with st.form("doc_contract"):
                client_sel = st.selectbox("العميل", options=clients, format_func=lambda x: x["name"])
                package = st.selectbox("الباقة", ["الأساسية", "القياسية", "الاحترافية"])
                price = st.text_input("السعر", value=settings.get("basic_price", "2,000"))
                duration = st.selectbox("المدة", ["شهر واحد", "3 أشهر", "6 أشهر"])
                start_date = st.date_input("تاريخ البدء")
                submitted = st.form_submit_button("📄 توليد العقد", type="primary")
                if submitted:
                    project_data = {
                        "name": f"مشروع {client_sel['name']}", "package": package,
                        "price": price, "start_date": start_date.strftime("%Y-%m-%d"),
                        "end_date": (start_date + timedelta(days=30)).strftime("%Y-%m-%d"),
                        "status": "قيد الانتظار", "scope": "إدارة حسابات السوشيال ميديا"
                    }
                    buf = generate_contract_doc(client_sel, project_data)
                    st.success("✅ تم إنشاء العقد!")
                    st.markdown(get_download_link(buf, f"عقد_{client_sel['name']}.docx", "📥 تحميل العقد (Word)"), unsafe_allow_html=True)
    with tab2:
        with st.form("doc_quotation"):
            q_client = st.text_input("اسم العميل", placeholder="اسم المطعم أو الكافيه")
            q_package = st.selectbox("الباقة", ["الأساسية", "القياسية", "الاحترافية"], key="q_pkg2")
            q_price = st.text_input("السعر", value=settings.get("basic_price", "2,000"), key="q_price2")
            q_notes = st.text_area("ملاحظات", placeholder="أي تفاصيل إضافية...")
            submitted = st.form_submit_button("💰 توليد عرض الأسعار", type="primary")
            if submitted:
                buf = generate_quotation_pdf(q_client, q_package, q_price, q_notes)
                st.success("✅ تم إنشاء عرض الأسعار!")
                st.markdown(get_download_link(buf, f"عرض_اسعار_{q_client}.pdf", "📥 تحميل عرض الأسعار (PDF)"), unsafe_allow_html=True)

elif menu == "⚙️ الإعدادات":
    st.markdown('<div class="header-gradient"><h1>⚙️ الإعدادات</h1></div>', unsafe_allow_html=True)
    with st.form("settings_form"):
        st.markdown("### 🏢 بيانات الشركة")
        company_name = st.text_input("اسم الشركة", value=settings.get("company_name", "محتوى برو"))
        phone = st.text_input("رقم الهاتف", value=settings.get("phone", ""))
        email = st.text_input("البريد الإلكتروني", value=settings.get("email", ""))
        st.markdown("---")
        st.markdown("### 💵 أسعار الباقات")
        basic = st.text_input("الحزمة الأساسية", value=settings.get("basic_price", "2,000 - 3,500"))
        standard = st.text_input("الحزمة القياسية", value=settings.get("standard_price", "4,500 - 8,000"))
        pro = st.text_input("الحزمة الاحترافية", value=settings.get("pro_price", "9,500 - 18,000"))
        st.markdown("---")
        st.markdown("### 🗑️ إدارة البيانات")
        col1, col2, col3 = st.columns(3)
        with col1:
            clear_clients = st.checkbox("حذف جميع العملاء")
        with col2:
            clear_projects = st.checkbox("حذف جميع المشاريع")
        with col3:
            clear_all = st.checkbox("حذف كل البيانات")
        submitted = st.form_submit_button("💾 حفظ الإعدادات", type="primary")
        if submitted:
            new_settings = {
                "company_name": company_name, "phone": phone, "email": email,
                "basic_price": basic, "standard_price": standard, "pro_price": pro
            }
            save_data("settings", new_settings)
            if clear_all:
                for key in ["clients", "projects", "tasks", "content", "payments"]:
                    save_data(key, [])
                st.warning("🗑️ تم حذف جميع البيانات!")
            else:
                if clear_clients:
                    save_data("clients", [])
                    st.warning("🗑️ تم حذف العملاء!")
                if clear_projects:
                    save_data("projects", [])
                    st.warning("🗑️ تم حذف المشاريع!")
            st.success("✅ تم حفظ الإعدادات")
            st.rerun()
    st.markdown("---")
    st.markdown("### 📤 تصدير البيانات")
    col1, col2, col3 = st.columns(3)
    with col1:
        clients_data = load_data("clients")
        if clients_data:
            csv = pd.DataFrame(clients_data).to_csv(index=False).encode('utf-8-sig')
            st.download_button("📥 تصدير العملاء", csv, "clients.csv", "text/csv")
    with col2:
        projects_data = load_data("projects")
        if projects_data:
            csv = pd.DataFrame(projects_data).to_csv(index=False).encode('utf-8-sig')
            st.download_button("📥 تصدير المشاريع", csv, "projects.csv", "text/csv")
    with col3:
        payments_data = load_data("payments")
        if payments_data:
            csv = pd.DataFrame(payments_data).to_csv(index=False).encode('utf-8-sig')
            st.download_button("📥 تصدير المالية", csv, "payments.csv", "text/csv")

st.markdown("---")
st.markdown(f"<p style='text-align:center;color:#666;'>🚀 {settings.get('company_name', 'محتوى برو')} - نظام إدارة المشاريع © 2026</p>", unsafe_allow_html=True)
